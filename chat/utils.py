from django.http import HttpResponse
from django.conf import settings
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def export_chat_to_docx(chat_session):
    """Export chat session to DOCX file"""
    try:
        # Create new document
        doc = DocxDocument()
        
        # Add title
        title = doc.add_heading(f'Chat: {chat_session.title}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f'Document: {chat_session.document.title}')
        doc.add_paragraph(f'Date: {chat_session.created_at.strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'User: {chat_session.user.username}')
        doc.add_paragraph('')  # Empty line
        
        # Add messages
        messages = chat_session.messages.all()
        for message in messages:
            if message.role == 'user':
                p = doc.add_paragraph()
                p.add_run('User: ').bold = True
                p.add_run(message.content)
            else:
                p = doc.add_paragraph()
                p.add_run('AI: ').bold = True
                p.add_run(message.content)
            doc.add_paragraph('')  # Empty line between messages
        
        # Save document
        export_dir = os.path.join(settings.MEDIA_ROOT, 'exports')
        os.makedirs(export_dir, exist_ok=True)
        
        filename = f"chat_{chat_session.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(export_dir, filename)
        doc.save(file_path)
        
        return f"/media/exports/{filename}"
        
    except Exception as e:
        raise Exception(f"Error exporting chat: {str(e)}")
